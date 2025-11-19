from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_summary_docx():
    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading('Summary of Organization and Management Chapters', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)

    # Chapter 1 Summary
    doc.add_heading('Chapter 1: Organization and Management', level=2)
    doc.add_paragraph(
        "Organization:\n"
        "- **Definition**: A social structure where people work interdependently to achieve common goals, coordinating human, material, and technological resources to create value for stakeholders.\n"
        "- **Characteristics**: Goal-oriented, social, and technological system, requiring cooperation, communication, and adherence to rules.\n"
        "- **Societal Impact**: Organizations (e.g., hospitals, universities, banks) enhance education, healthcare, and technology.\n"
        "- **System Approach**: Organizations are interdependent systems that adapt to external changes (e.g., technology, customer needs, regulations) using feedback.\n\n"
        "Principles of Organization (Henry Fayol):\n"
        "- **Division of Work**: Subdivide tasks for efficiency.\n"
        "- **Authority and Responsibility**: Managers have decision-making rights and obligations.\n"
        "- **Span of Control**: Defines the number of subordinates per supervisor (e.g., 20 employees per supervisor).\n"
        "- **Chain of Command**: Clear authority flow from top to bottom.\n"
        "- **Unity of Command**: Each employee reports to one superior.\n\n"
        "Management:\n"
        "- **Definition**: The act of coordinating people and resources to achieve organizational goals.\n"
        "- **Historical Context**: Concepts trace back to Sun Tzu (6th century BC) and Chanakya’s Arthashastra (300 BC).\n"
        "- **Levels**:\n"
        "  - **Top-Level**: Sets goals, policies, budgets, and makes decisions.\n"
        "  - **Middle-Level**: Implements policies as mediators between top and lower levels.\n"
        "  - **Lower-Level**: Supervises daily activities and workers.\n"
        "- **Functions**: Planning, organizing, directing, coordinating, controlling.\n"
        "- **Skills**: Interpersonal (interaction), informational (data gathering), decisional (decision-making).\n"
        "- **Models**:\n"
        "  - **Allocational**: Manages resources (money, manpower).\n"
        "  - **Team Effort**: Emphasizes cooperation and leadership.\n"
        "  - **Knowledge-Oriented**: Uses experience and training to adapt.\n\n"
        "Theories of Management:\n"
        "- **Scientific Management (Frederick Taylor)**: Emphasizes selecting suitable employees, training, and cooperation, increasing output (e.g., from 12.7 to 48.8 tons/day).\n"
        "- **Human Relations Movement (Elton Mayo)**: Hawthorne Experiments showed productivity improves with managerial attention and social factors.\n"
        "- **Administrative Management (Fayol)**: Views management as a scientific process with 14 principles (e.g., specialization, discipline).\n"
        "- **Contingency Approach**: Integrates theories, adapting to situational needs.\n\n"
        "Types of Organizations:\n"
        "- **Single Ownership**: Easy to form/dissolve, full control, but limited resources and unlimited liability.\n"
        "- **Partnership**: Combines skills, quick decisions, but risks conflicts.\n"
        "- **Joint Stock**: Large financial resources, limited liability, but complex formation and lack of secrecy.\n"
        "- **Cooperative Societies**: Democratic, government support, but limited capital and managerial talent.\n"
        "- **Public Corporation**: Unlimited resources, continuity, but excessive regulations and political influence.\n\n"
        "Organizational Structure:\n"
        "- **Line Organization**: Simple, clear authority, but top-heavy and unsuitable for large firms.\n"
        "- **Line and Staff**: Adds specialists to advise, but risks coordination issues.\n"
        "- **Functional Organization**: Divides work by functions (e.g., production, marketing), promoting specialization but complex.\n"
        "- **Committee Organization**: Pools expertise, but may delay decisions due to compromises.\n\n"
        "Purchasing and Marketing:\n"
        "- **Purchasing**: Ensures timely material supply via direct purchase, quotations, or tenders. Steps include identifying needs, evaluating suppliers, and inspecting materials.\n"
        "- **Marketing**: Plans, prices, promotes, and distributes products/services to satisfy customers and achieve profits. Functions include advertising and customer-oriented strategies."
    )

    # Chapter 2 Summary
    doc.add_heading('Chapter 2: Personnel Management', level=2)
    doc.add_paragraph(
        "Personnel Management Overview:\n"
        "- Focuses on managing human resources to achieve organizational goals through planning, hiring, development, and motivation.\n\n"
        "Aspects and Functions:\n"
        "- **Managerial Functions**: Planning, organizing, directing, controlling.\n"
        "- **Operative Functions**:\n"
        "  - **Procurement/Employment**: Hiring the right number/type of employees via recruitment, selection, placement.\n"
        "  - **Development**: Enhances skills through training, seminars, and team-building.\n"
        "  - **Compensation**: Sets fair wages based on job requirements, employee needs, market rates, and organizational capacity.\n"
        "  - **Integration**: Maintains communication and harmony between employees, employers, and trade unions.\n"
        "  - **Maintenance**: Ensures health, safety, and favorable working conditions.\n"
        "  - **Motivation**: Uses monetary (e.g., bonuses) and non-monetary (e.g., flexible schedules, social activities) incentives.\n\n"
        "Manpower Planning:\n"
        "- Enhances productivity, reduces labor excess/shortages, attracts talent, and lowers operating costs.\n\n"
        "Job Analysis:\n"
        "- Defines tasks, duties, skills, tools, and compensation to improve workflow.\n"
        "- **Components**:\n"
        "  - **Job Description**: Details job title, responsibilities, location, conditions, salary.\n"
        "  - **Job Specification**: Lists required qualifications, skills, experience.\n"
        "  - **Job Design**: Modifies job structure for productivity.\n"
        "  - **Job Evaluation**: Assesses job value for fair compensation.\n"
        "- **Methods**: Questionnaire (job complexity), interview (detailed duties), observation (working conditions, but behavior may change).\n\n"
        "Recruitment and Selection:\n"
        "- **Steps**: Job analysis, advertising vacancies, application collection, reference checks, tests, selection, induction (introducing policies), placement.\n\n"
        "Training and Development:\n"
        "- Continuous training boosts skills, motivation, morale, and efficiency, reducing turnover and supervision needs.\n\n"
        "Wage/Salary Structure:\n"
        "- Influenced by organizational capacity, labor market supply/demand, market rates, living wage, productivity, and trade union bargaining.\n\n"
        "Incentives:\n"
        "- **Monetary**: Profit-sharing, bonuses, salary increases.\n"
        "- **Non-Monetary**: Flexible schedules, holidays, social activities, gift vouchers.\n"
        "- **Individual**: Promotions, improved conditions for personal contributions.\n"
        "- **Group**: Rewards team output, encouraging collaboration.\n\n"
        "Performance Appraisal:\n"
        "- Provides feedback, informs promotions, salary reviews, job rotations, and training.\n"
        "- **Methods**:\n"
        "  - **Traditional**: Evaluates traits (e.g., creativity, responsibility).\n"
        "  - **Modern**: Management by Objectives (MBO) sets goals and assesses results; Assessment Center evaluates interpersonal skills and potential."
    )

    # Chapter 3 Summary
    doc.add_heading('Chapter 3: Motivating and Leading People', level=2)
    doc.add_paragraph(
        "Motivation Overview:\n"
        "- **Definition**: Influencing action by aligning organizational and personal goals, driven by intrinsic (internal desire) and extrinsic (external rewards/punishment) factors.\n"
        "- **Intrinsic vs. Extrinsic**: Intrinsic is self-driven (e.g., enjoyment); extrinsic involves rewards (e.g., bonuses) or fear of punishment.\n"
        "- **Positive vs. Negative**: Rewards for tasks vs. consequences for failure.\n\n"
        "Role of Management:\n"
        "- Creates motivating environments through training, clear goals, feedback, fair wages, and fun workplaces.\n"
        "- **Techniques**: Offer meaningful/challenging work, set expectations, provide supportive feedback, value self-motivated employees.\n\n"
        "Types of Motivation:\n"
        "- **Attitude**: Enhances performance via self-confidence and positive outlook.\n"
        "- **Group**: Boosts team performance with incentives and recognition.\n"
        "- **Executive**: Retains talent with profit-sharing, challenging roles, and benefits.\n\n"
        "Theories of Motivation:\n"
        "- **Content Theories**: Maslow’s Hierarchy, Alderfer’s ERG, McClelland’s Needs, Herzberg’s Two-Factor (hygiene factors prevent dissatisfaction; motivators drive satisfaction).\n"
        "- **Process Theories**: Vroom’s Expectancy Theory.\n"
        "- **McGregor’s Theory X/Y**: X assumes workers are lazy, need control; Y assumes they are self-motivated.\n\n"
        "Leadership:\n"
        "- **Definition**: Guiding people toward goals via inspiration and resource allocation.\n"
        "- **Leaders vs. Managers**: Managers plan/control; leaders inspire. Effective managers are leaders.\n"
        "- **Qualities**: Trustworthiness, honesty, confidence, communication, discipline, imagination, commitment.\n"
        "- **Styles (Managerial Grid)**:\n"
        "  - **Impoverished (1,1)**: Low concern for people/production.\n"
        "  - **Country Club (1,9)**: High people concern, low production.\n"
        "  - **Dictatorial (9,1)**: High production, low people concern.\n"
        "  - **Team (9,9)**: High concern for both, fostering teamwork.\n"
        "  - **Middle-of-the-Road (5,5)**: Balances both moderately.\n"
        "- **General Styles**:\n"
        "  - **Autocratic**: Unilateral decisions, quick but may cause resentment.\n"
        "  - **Democratic**: Group input, boosts productivity/morale.\n"
        "  - **Laissez-Faire**: Minimal direction, suitable for skilled teams.\n"
        "- **Approaches**:\n"
        "  - **Trait**: Leaders have qualities like integrity, intelligence.\n"
        "  - **Behavioral**: Leaders are made, focusing on people-oriented (trust) or task-oriented (goals) behaviors.\n"
        "  - **Contingency/Situational**: Adapts to situations with directive, supportive, participative, or achievement-oriented styles.\n"
        "  - **Integrated**: Combines traits, behaviors, follower maturity, and context.\n\n"
        "Entrepreneurship:\n"
        "- **Definition**: Starting/managing ventures with risks for profit, driven by creativity and initiative.\n"
        "- **Characteristics**: Opportunity-seeking, calculated risk-taking, creativity, strong work ethic.\n"
        "- **Role**: Drives economic growth via innovation, job creation, and new markets.\n"
        "- **Promotion**: Governments foster entrepreneurial culture (e.g., Global Entrepreneurship Week).\n"
        "- **Venture Setup**: Choose ownership (sole, partnership), location, approvals, financing, production, marketing."
    )

    # Chapter 5 Summary
    doc.add_heading('Chapter 5: Management Information System', level=2)
    doc.add_paragraph(
        "MIS Overview:\n"
        "- **Definition**: Uses computing and communication technology to collect, analyze, and organize information for efficient decision-making.\n"
        "- **Purpose**: Supports planning, organizing, directing, coordinating, and controlling with insights into past, present, and future activities.\n\n"
        "Key Components:\n"
        "- **Management, Information, System**.\n"
        "- **Data vs. Information**: Data are raw facts; information is analyzed, relevant data.\n\n"
        "Functions of MIS:\n"
        "- **Data Capturing**: Collects internal/external data.\n"
        "- **Processing**: Converts data into useful information.\n"
        "- **Prediction**: Forecasts future scenarios using statistics.\n"
        "- **Planning**: Sets goals and targets.\n"
        "- **Controlling**: Identifies problems and actions.\n"
        "- **Assistance**: Provides timely, accurate information for decisions.\n\n"
        "Evolution:\n"
        "- Evolved from manual systems (e.g., punch cards) due to management theories, organizational changes, and computer technology.\n\n"
        "MIS Across Management Levels:\n"
        "- **Strategic**: Long-term planning, forecasting, aggregated data (periodic).\n"
        "- **Tactical**: Short-term budgeting, inventory, personnel (weekly/monthly).\n"
        "- **Operational**: Daily production, real-time data for scheduling (hourly/daily).\n\n"
        "Role of Computers:\n"
        "- Overcome pre-computer issues (late, incomplete, costly information).\n"
        "- Support online, real-time, and time-sharing processing.\n\n"
        "Database:\n"
        "- **Components**: Data dictionary, architecture, manipulation rules.\n"
        "- Organizes data for efficient retrieval.\n\n"
        "Networking:\n"
        "- Enables interconnected data transmission and integrated systems.\n\n"
        "Information System Model:\n"
        "- Involves users, business sources, programming, processing, feedback.\n\n"
        "Information Support:\n"
        "- **Planning**: Supports forecasting and resource utilization.\n"
        "- **Decision-Making**: Strategic (long-term), tactical (short-term).\n"
        "- **Information Architecture**: Organizes content for user efficiency."
    )

    # Differentiation of Management Levels and MIS Needs
    doc.add_heading('Differentiation of Management Levels and MIS Needs', level=2)
    doc.add_paragraph(
        "Management levels differ in their focus, decision types, and MIS requirements, as outlined below:"
    )

    # Create a table for differentiation
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Management Level', 'Focus', 'Decision Type', 'MIS Needs', 'Frequency of Information', 'Examples of MIS Use']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Data rows
    data = [
        ('Strategic', 'Long-term goals, policies, budgets', 'Future-oriented, high-level', 'Forecasting, aggregated data, trend analysis', 'Periodic (quarterly, annually)', 'Market trend reports, budget planning'),
        ('Tactical', 'Short-term plans, resource allocation', 'Implementation-focused, departmental', 'Budget tracking, KPIs, resource data', 'Frequent (weekly, monthly)', 'Inventory schedules, personnel reports'),
        ('Operational', 'Daily production, task execution', 'Immediate, task-specific', 'Real-time data, production monitoring', 'Very frequent (hourly, daily)', 'Work schedules, raw material tracking'),
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    # Save the document
    doc.save('Management_Summaries.docx')

if __name__ == '__main__':
    create_summary_docx()